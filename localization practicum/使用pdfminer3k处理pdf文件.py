#!/usr/bin/env python
# coding: utf-8

# 使用pdfminer3k处理pdf格式文件
# + 安装pdfminer3k：pip3 install pdfminer3k
# + pdfminer3k文档
#   * http://denis.papathanasiou.org/archive/2010.08.04.post.pdf
#   * http://www.unixuser.org/~euske/python/pdfminer/index.html
# + pdfminder code-https://github.com/dpapathanasiou/pdfminer-layout-scanner

# In[103]:


# 引用部分：
from pdfminer.pdfparser import PDFParser, PDFDocument, PDFNoOutlines
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine, LTFigure, LTImage


# In[80]:


#基本使用方法
#1、打开PDF文件
fp = open('D:/BLCU/大四上/行业综述论文/Definition.pdf', 'rb') #buffer
#print(fp)


# In[81]:


#2、创建与文件对象相关联的PDF解释对象-Create a PDF parser object associated with the file object.
parser = PDFParser(fp)
#print(parser)    


# In[91]:


#3、创建一个PDF文件对象，用于存储文件结构，如果pdf有密码就填写密码，没有就空着，为初始化pdf文件做准备。
# Create a PDF document object that stores the document structure.
# Supply the password for initialization. -document = PDFDocument(parser, password)
document = PDFDocument(parser)
#print(document) = <pdfminer.pdfparser.PDFDocument object at 0x064AF7D8>
document.initialize("") #原文档没有写，但不初始化就会说没有功能
print(document)

doc = PDFDocument()
#连接分析器，与文档对象
parser.set_document(doc)
doc.set_parser(parser)


# In[83]:


# 4、检查所创建的文件对象是否允许文本提取，如果不允许提取文本，则终止。Check if the document allows text extraction. If not, abort.
if not document.is_extractable:
    raise PDFTextExtractionNotAllowed
    
else:
    print('可以读取')


# In[84]:


#5、创建PDF资源管理对象以村塾共享资源-# Create a PDF resource manager object that stores shared resources.
rsrcmgr = PDFResourceManager()
print(rsrcmgr)


# In[85]:


#6、创建PDF device对象
from pdfminer.pdfdevice import PDFDevice
device = PDFDevice(rsrcmgr)
print(device)


# In[86]:


# 7、创建一个PDF解释对象Create a PDF interpreter object. （manager和device作为参数
interpreter = PDFPageInterpreter(rsrcmgr, device)
print(interpreter)


# In[87]:


# 8、处理PDF文件中的每一页面内容-Process each page contained in the document.
for page in PDFPage.create_pages(document): #PDFPage库在PDFMiner3k中已经不适用了，需要替换
    interpreter.process_page(page)


# In[92]:


# 9、循环遍历列表，每次处理一页的内容
# doc.get_pages() 获取page列表
for page in doc.get_pages():
    # 利用解释器的process_page()方法解析读取单独页数
    interpreter.process_page(page)


# In[104]:


from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator

# Set parameters for analysis.
laparams = LAParams()
# Create a PDF page aggregator object.
device = PDFPageAggregator(rsrcmgr, laparams=laparams)
interpreter = PDFPageInterpreter(rsrcmgr, device)
for page in PDFPage.create_pages(document):
    interpreter.process_page(page)
    # receive the LTPage object for the page.
    layout = device.get_result()


# In[108]:


#10、获取目录
from pdfminer.pdfparser import PDFParser

# Open a PDF document.
fp = open('foriegn.pdf', 'rb')
parser = PDFParser(fp)
document = PDFDocument(parser)

# Get the outlines of the document.
outlines = document.get_outlines()
for (level,title,dest,a,se) in outlines:
    print (level, title)


# 接下来是完整版本代码实现：

# In[27]:


#代码-PDF转word或txt
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
 
document = Document()

def parse():
    # rb以二进制读模式打开本地pdf文件
    fn = open('D:/BLCU/大四上/行业综述论文/Definition.pdf','rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
 
    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
 
    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource,laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource,device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out,"get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    with open('test.txt','a') as f:  #存储txt文本
                        f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'    # 添加段落，样式为unordered list类型
                    )
                document.save('demo.docx')  # 保存这个文档
 
 
if __name__ == '__main__':
    parse()


# In[ ]:





# In[109]:


import pyocr
import importlib
import sys
import time
 
importlib.reload(sys)
time1 = time.time()
# print("初始时间为：",time1)
 
import os.path
from pdfminer.pdfparser import  PDFParser,PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal,LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
 
text_path = r'foriegn.pdf'
# text_path = r'photo-words.pdf'
 
def parse():
    '''解析PDF文本，并保存到TXT文件中'''
    fp = open(text_path,'rb')
    #用文件对象创建一个PDF文档分析器
    parser = PDFParser(fp)
    #创建一个PDF文档
    doc = PDFDocument()
    #连接分析器，与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
 
    #提供初始化密码，如果没有密码，就创建一个空的字符串
    doc.initialize()
 
    #检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        #创建PDF，资源管理器，来共享资源
        rsrcmgr = PDFResourceManager()
        #创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr,laparams=laparams)
        #创建一个PDF解释其对象
        interpreter = PDFPageInterpreter(rsrcmgr,device)
 
        #循环遍历列表，每次处理一个page内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            interpreter.process_page(page)
            #接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
            # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
            # 想要获取文本就获得对象的text属性，
            for x in layout:
                if(isinstance(x,LTTextBoxHorizontal)):
                    with open(r'post.txt','a',encoding='utf-8') as f: #gbk有的字符解析不了，所以加了个encoding=utf-8
                        results = x.get_text()
                        print(results)
                        f.write(results  +"\n")


if __name__ == '__main__':
    parse()
    time2 = time.time()
    print("总共消耗时间为:",time2-time1)


# In[73]:


import urllib
import importlib,sys
importlib.reload(sys)
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
 
def parse(DataIO, save_path):
 
    #用文件对象创建一个PDF文档分析器
    parser = PDFParser(DataIO)
    #创建一个PDF文档
    doc = PDFDocument()
    #分析器和文档相互连接
    parser.set_document(doc)
    doc.set_parser(parser)
    #提供初始化密码，没有默认为空
    doc.initialize()
    #检查文档是否可以转成TXT，如果不可以就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
    else:
        #创建PDF资源管理器，来管理共享资源
        rsrcmagr = PDFResourceManager()
        #创建一个PDF设备对象
        laparams = LAParams()
        #将资源管理器和设备对象聚合
        device = PDFPageAggregator(rsrcmagr, laparams=laparams)
        #创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmagr, device)
 
        #循环遍历列表，每次处理一个page内容
        #doc.get_pages()获取page列表
        for page in doc.get_pages():
            interpreter.process_page(page)
            #接收该页面的LTPage对象
            layout = device.get_result()
            #这里的layout是一个LTPage对象 里面存放着page解析出来的各种对象
            #一般包括LTTextBox，LTFigure，LTImage，LTTextBoxHorizontal等等一些对像
            #想要获取文本就得获取对象的text属性
            for x in layout:
                try:
                    if(isinstance(x, LTTextBoxHorizontal)):
                        with open('%s' % (save_path), 'a') as f:
                            result = x.get_text()
                            print (result)
                            f.write(result + "\n")
                except:
                    print("Failed")
 
 
if __name__ == '__main__':
    #解析本地PDF文本，保存到本地TXT
    with open(r'Definition.pdf','rb') as pdf_html:
        parse(pdf_html, r'test.txt')
 
    #解析网络上的PDF，保存文本到本地
    # url = "https:"
    # pdf_html = urllib.urlopen(url).read()
    # DataIO = StringIO(pdf_html)
    # parse_pdf(DataIO, r'E:\parse_pdf')

